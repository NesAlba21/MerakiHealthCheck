from docx import Document
from docx.shared import Pt
from MerakiAPI import apiInit
from GET_orgs import getOrgs
from GET_NetworksBind import getNetworksBind
from GET_Templates import getTemplate
from GET_firmware import firmware
from GET_License import getLicence
from Get_WarmSpare import getWarmSpare
from GET_AlertsNetwork import getAlertsNetwork
from GET_NetworkDevices import getNetworkDevices
from Get_SwitchStacksNetwork import getSwitchStack
from GET_MS_mtu import mtu
from GET_MsDetails import getMsDetail
from GET_VpnBgp import getVpnBgp
from GET_Site2Site_VPN import getNetworkVpn
from GET_MxSetting import  getMxSettings
from GET_MxVlans import getNetMxVlan
from GET_snmp import getSnmp
from GET_Ssid import getSsid
from GET_rfProfiles import get_rfProfiles
from GET_l3Fw import getl3Fw

api = apiInit()
# Put the output in a Word document using python-docx"""
# Create Document
document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(8)



# Add Title
document.add_heading('Meraki Health Check', 0)
document.add_paragraph(' ')

def addSection(table):
    # Add Header before the table
    document.add_heading(table.title,4)
    document.add_paragraph(description)


    # Add a table to the document
    rows = len(table._rows)+1
    cols = len(table.field_names)
    table_word = document.add_table(rows=rows, cols=cols, style='Light List Accent 1')



    # Add header row to the table
    for i, field in enumerate(table.field_names):
        cell = table_word.cell(0, i)
        cell.text = field

    # Add data rows to the table
    for i, row in enumerate(table._rows):
        for j, value in enumerate(row):
            cell = table_word.cell(i + 1, j)
            cell.text = str(value)

    document.add_paragraph(' ')

document.add_heading("Organization Level", 1)


# Your Organizations
orgId, table = getOrgs(api)
description = ("Below you can find a list of organizations in which you have privileges")
addSection(table)

# Organization Networks
network_id, selected_option, table = getNetworksBind(api, orgId)
desciption = ("Below you can find the Networks inside your Organization")
addSection(table)

# Organization Licenses
table = getLicence(api,orgId)
description = ("Below you can find the Organization licenses")
addSection(table)

# Organization Templates
table = getTemplate(api, orgId)
description = ("Below you can find the configured Organizations Templates ")
addSection(table)

# SNMP
table = getSnmp(api, orgId)
description = ("Organization snmp settings")
addSection(table)



document.add_heading("Network Level", 1)
document.add_heading("Devices", 2)
# Network Devices
list_serial, table = getNetworkDevices(api, network_id, selected_option)
description = "Below you can find the devices list  in a network "
addSection(table)

# Firmware
table = firmware(api, network_id)
description = ("Below you can find the current firmware version per device")
addSection(table)

document.add_heading("Alerts", 2)
# Network Alerts
table1, table = getAlertsNetwork(api, network_id)
description = ("Default alerts configuration on the network")
addSection(table1)
description = ("Default individual alerts configuration on the network")
addSection(table)

document.add_heading("MS(switches) info", 2)
# Network Switches info
table1, table = getMsDetail(api, list_serial)
description = ('MS switch port statuses')
addSection(table1)
description = ("MS port details")
addSection(table)

# Network switch stack info
table = getSwitchStack(api, network_id)
description = "Below you can find the swtich stack information on the network"
addSection(table)

# Network MTU configuration
table = mtu(api, network_id)
description = "Recommended to keep at default of 9578 unless intermediate devices don’t support jumbo frames. " \
              "This is useful to optimize server-to-server and application performance. Avoid fragmentation when possible."
addSection(table)

document.add_heading("MX(appliance) info", 2)

# Network MX settings
table = getMxSettings(api, network_id)
description = ("The Cisco Meraki MX security appliance has a number of deployment options to meet the needs "
               "of your network and infrastructure. Whether as the main edge firewall for your network, or as a concentrator device in your data center, the MX security appliance can be easily integrated.")
addSection(table)

# Network MX Vlans
table = getNetMxVlan(api, network_id)
description = ("List the VLANs for an MX network")
addSection(table)

# Network L3 Firewall Rules
table = getl3Fw(api, network_id)
description = ("Layer 3 Firewall rules configured on the network")
addSection(table)

# Network MX in warmspare
table = getWarmSpare(api, network_id)
description = ("Below you can find the MX Warmspare details on the network")
addSection(table)

document.add_heading("BGP", 3)
# Network MX Hub BGP Configuration
table1, table = getVpnBgp(api, network_id)
p = document.add_paragraph("** Deployment mode should be configured as ")
p.add_run('passthrough').bold = True

description = ( "MX appliances use iBGP to exchange route information over Meraki AutoVPN. "
               "MXs deployed as one-armed VPN concentrator use eBGP to exchange and learn routes within the datacenter."
               " Learned routes are redistributed to AutoVPN spokes using iBGP.")
addSection(table)
addSection(table1)

document.add_heading("Site-to-Site VPN", 3)
# Network site-to-site VPN settings
table3, table2, table = getNetworkVpn(api, network_id)
description = "Meraki Auto VPN technology is a unique solution that allows site-to-site VPN tunnel creation"
addSection(table)
addSection(table2)
addSection(table3)


document.add_heading("MR(wireless) info", 2)

# ssid's
table = getSsid(api, network_id)
description = "List of SSIDs in the network"
addSection(table)

# RF profiles
table = get_rfProfiles(api, network_id)
description = "List of wireless Radio Frequency profiles in the network"
addSection(table)


document.add_page_break()



# Save the document
document.save('Meraki Health Check.docx')
print("")
print("Document successfully created!")
